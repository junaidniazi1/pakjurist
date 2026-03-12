"""
PakJurist - AI Agent Module
Handles all AI interactions, document processing, voice recognition, and text-to-speech.

INSTALLATION (Windows):
    pip install torch torchaudio --index-url https://download.pytorch.org/whl/cpu
    pip install coqui-tts

NOTE: Do NOT use `pip install TTS` — that package is abandoned and broken on Windows.
      Use `pip install coqui-tts` (the actively maintained fork).
"""

import google.generativeai as genai
import os
import io
import base64
from PIL import Image
import PyPDF2
import docx
import speech_recognition as sr
import tempfile
from typing import Optional, List, Union
from dataclasses import dataclass
from datetime import datetime


@dataclass
class ProcessedDocument:
    """Data class for processed documents"""
    name: str
    doc_type: str
    content: Union[str, Image.Image]
    processed_at: str


@dataclass
class TTSResult:
    """Data class for TTS output"""
    audio_bytes: bytes       # raw WAV bytes
    audio_b64: str           # base64-encoded WAV for embedding in HTML <audio>
    sample_rate: int
    language: str
    model_used: str
    generated_at: str


class CoquiTTSEngine:
    """
    Advanced Text-to-Speech engine powered by Coqui TTS (coqui-tts fork).

    Features:
    - Multilingual voices (English, Urdu, and many more)
    - Voice cloning via XTTS v2 (provide a reference WAV ≥6 seconds)
    - Realistic, natural-sounding speech
    - Fully open-source — no API key required
    - Works on Windows, Linux, and macOS

    Install:
        pip install torch torchaudio --index-url https://download.pytorch.org/whl/cpu
        pip install coqui-tts

    Models are lazy-loaded on first use and cached locally by Coqui.
    First call will trigger a one-time model download.
    """

    # Best single-language model per supported language key
    LANGUAGE_MODEL_MAP = {
        'english': 'tts_models/en/ljspeech/tacotron2-DDC',
        'urdu':    'tts_models/ur/cv/vits',                  # Native Urdu VITS
        'roman':   'tts_models/en/ljspeech/tacotron2-DDC',   # Roman Urdu → English model
    }

    # XTTS v2: high-quality multilingual model with voice cloning support
    MULTILINGUAL_MODEL = 'tts_models/multilingual/multi-dataset/xtts_v2'

    def __init__(self, use_multilingual: bool = False, use_gpu: bool = False):
        """
        Args:
            use_multilingual: Load XTTS v2 for better quality + voice cloning.
                              Requires ~4 GB RAM. Recommended for production.
            use_gpu:          Use CUDA GPU for faster inference.
        """
        self._tts = None
        self._current_model: Optional[str] = None
        self.use_multilingual = use_multilingual
        self.use_gpu = use_gpu

    def _load_model(self, model_name: str) -> None:
        """Lazy-load a Coqui TTS model (downloads on first use, then cached)."""
        try:
            from TTS.api import TTS  # noqa: internal import name stays "TTS" inside coqui-tts
        except ImportError:
            raise ImportError(
                "Coqui TTS is not installed.\n"
                "Run the following commands:\n"
                "  pip install torch torchaudio --index-url https://download.pytorch.org/whl/cpu\n"
                "  pip install coqui-tts\n"
                "Do NOT use `pip install TTS` — that package is abandoned and broken on Windows."
            )
        if self._current_model != model_name:
            self._tts = TTS(model_name=model_name, gpu=self.use_gpu)
            self._current_model = model_name

    def _resolve_model(self, language: str) -> str:
        """Return the correct model name for a given language key."""
        if self.use_multilingual:
            return self.MULTILINGUAL_MODEL
        return self.LANGUAGE_MODEL_MAP.get(language, self.LANGUAGE_MODEL_MAP['english'])

    # Language key → BCP-47 code (used by multilingual models)
    _LANG_CODE = {'english': 'en', 'urdu': 'ur', 'roman': 'en'}

    def synthesize(
        self,
        text: str,
        language: str = 'english',
        speaker_wav: Optional[str] = None,
        speaker: Optional[str] = None,
    ) -> TTSResult:
        """
        Synthesize speech from text and return WAV bytes + base64.

        Args:
            text:        Text to convert to speech.
            language:    'english', 'urdu', or 'roman'.
            speaker_wav: Path to a reference WAV for voice cloning (XTTS v2 only, ≥6 s).
            speaker:     Named speaker ID for multi-speaker models.

        Returns:
            TTSResult with raw WAV bytes, base64 string, and metadata.
        """
        model_name = self._resolve_model(language)
        self._load_model(model_name)
        lang_code = self._LANG_CODE.get(language, 'en')

        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            kwargs: dict = {"text": text, "file_path": tmp_path}

            if self.use_multilingual:
                kwargs["language"] = lang_code
                if speaker_wav:
                    kwargs["speaker_wav"] = speaker_wav
                elif speaker:
                    kwargs["speaker"] = speaker
            else:
                if speaker:
                    kwargs["speaker"] = speaker

            self._tts.tts_to_file(**kwargs)

            with open(tmp_path, "rb") as f:
                audio_bytes = f.read()

            import wave
            with wave.open(tmp_path, 'rb') as wf:
                sample_rate = wf.getframerate()

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

        audio_b64 = base64.b64encode(audio_bytes).decode('utf-8')

        return TTSResult(
            audio_bytes=audio_bytes,
            audio_b64=audio_b64,
            sample_rate=sample_rate,
            language=language,
            model_used=model_name,
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        )

    def clone_voice_and_speak(
        self,
        text: str,
        speaker_wav_path: str,
        language: str = 'english',
    ) -> TTSResult:
        """
        Synthesize speech in a cloned voice using XTTS v2.

        Args:
            text:             Text to synthesize.
            speaker_wav_path: Path to a clean reference WAV (≥6 seconds recommended).
            language:         Target language.

        Returns:
            TTSResult with cloned-voice audio.

        Raises:
            ValueError: If the multilingual model is not enabled.
        """
        if not self.use_multilingual:
            raise ValueError(
                "Voice cloning requires the multilingual XTTS v2 model. "
                "Initialize CoquiTTSEngine(use_multilingual=True)."
            )
        return self.synthesize(text, language=language, speaker_wav=speaker_wav_path)

    def list_available_models(self) -> List[str]:
        """List all models available in the installed coqui-tts package."""
        try:
            from TTS.api import TTS
            return TTS().list_models()
        except ImportError:
            raise ImportError(
                "coqui-tts is not installed. Run:\n"
                "  pip install torch torchaudio --index-url https://download.pytorch.org/whl/cpu\n"
                "  pip install coqui-tts"
            )

    def list_speakers(self, language: str = 'english') -> List[str]:
        """
        Return named speaker IDs for the active model.
        Returns an empty list for single-speaker models.
        """
        self._load_model(self._resolve_model(language))
        if self._tts and getattr(self._tts, 'speakers', None):
            return list(self._tts.speakers)
        return []


class LegalAgent:
    """
    AI Agent for PakJurist — Pakistan Legal Awareness System.
    Provides legal information, document analysis, speech-to-text input,
    and advanced text-to-speech output via Coqui TTS.
    """

    SYSTEM_PROMPT = """You are PakJurist, an AI-powered Pakistan Legal Awareness and Constitutional Assistance System. Your purpose is to provide accurate, structured, and educational information about the laws of Pakistan.

Your responses may cover:
- Constitutional law of Pakistan
- Criminal law (e.g., Pakistan Penal Code, Criminal Procedure Code)
- Civil law
- Commercial and business law
- Administrative regulations
- General legal principles applicable within Pakistan

Your primary users include tourists visiting Pakistan, citizens seeking legal awareness, law students completing academic tasks, and legal professionals who need structured explanations of established legal provisions.

---

CRITICAL RULES:

1. NO HALLUCINATION
   Never invent laws, articles, sections, punishments, amendments, procedures, or court decisions.

2. ONLY ESTABLISHED LEGAL REFERENCES
   Mention only well-known and widely recognized provisions within Pakistani law.

3. UNCERTAIN INFORMATION
   If you are not confident that information is current or verified, clearly state:
   "I do not have confirmed latest information on this matter."

4. EDUCATIONAL PURPOSE ONLY
   Your responses are for educational and informational purposes only.
   You are not a lawyer and do not provide legal advice.

5. NO LEGAL STRATEGY OR CASE OUTCOME PREDICTIONS
   Do not speculate about case outcomes, legal strategies, or personal legal advice.

6. PROCEDURAL UNCERTAINTY
   If exact legal procedures, recent amendments, judicial precedents, or official interpretations cannot be verified, politely explain this limitation.

---

LANGUAGE RULES (STRICT):

- Default: Respond in the same language used by the user.
- If the user writes in Urdu script → reply in Urdu.
- If the user writes in English → reply in English.
- If the user writes in Roman Urdu → reply in Roman Urdu (you may add a short English explanation if helpful).
- If the user explicitly requests a language (e.g., "Urdu mein batao", "Reply in English") → follow that instruction.
- If the message includes instruction tags such as:
    [RESPOND IN URDU]
    [RESPOND IN ROMAN URDU AND ENGLISH]
  You MUST follow the instruction exactly.

---

SCOPE LIMITATION:

If a question is not related to Pakistani law, respond with:
"This question is not related to Pakistani law and no legal information is available."

---

DOCUMENT ANALYSIS:

When users upload images, PDFs, or documents:
- Carefully analyze the document for legal content.
- Extract relevant legal information such as case details, legal questions, or referenced laws.
- Provide responses based only on the document content and established Pakistani law.
- Do not assume missing facts from documents.

---

TRANSPARENCY:

- If information is missing, outdated, or uncertain, clearly state this instead of guessing.
- If asked how many laws exist in Pakistan, explain that Pakistan does not have a fixed number of laws because laws are continuously enacted, amended, repealed, and interpreted at both federal and provincial levels.
- When discussing amendments, only mention widely known and verified amendments. If verification is uncertain, clearly state that confirmation is required from an official legal source or government gazette.

---

SOURCE REFERENCES:

When your answer relies on a known legal source, mention the source clearly.

Whenever possible, include a reference link to an official or authoritative source such as:
- Government of Pakistan websites
- Official gazettes
- Court websites
- Recognized legal databases
- Educational or institutional legal resources

Always place the reference at the end of the explanation under a section titled:

"Reference Source"

Example format:

Reference Source:
- [Name of Law / Article / Court / Institution]
- Link: (official website or reliable legal source)

If a reliable reference link cannot be confirmed, clearly state:
"An official reference link is not currently available."

---

Always prioritize accuracy over completeness."""

    def __init__(
        self,
        api_key: Optional[str] = None,
        enable_tts: bool = True,
        tts_multilingual: bool = False,
        tts_gpu: bool = False,
    ):
        """
        Initialize PakJurist Legal Agent.

        Args:
            api_key:           Google Gemini API key (falls back to GEMINI_API_KEY env var).
            enable_tts:        Enable Coqui TTS engine for voice output.
            tts_multilingual:  Use XTTS v2 multilingual model (higher quality + voice cloning).
            tts_gpu:           Use GPU for TTS inference (faster on CUDA-enabled hardware).
        """
        self.api_key = api_key or os.getenv('GEMINI_API_KEY')
        if not self.api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")

        self.model = self._initialize_model()
        self.chat_session = None
        self.recognizer = sr.Recognizer()

        # Coqui TTS — lazily loads its model on first synthesize() call
        self.tts_engine: Optional[CoquiTTSEngine] = (
            CoquiTTSEngine(use_multilingual=tts_multilingual, use_gpu=tts_gpu)
            if enable_tts else None
        )

    # ------------------------------------------------------------------
    # Gemini model initialisation
    # ------------------------------------------------------------------

    def _initialize_model(self):
        """Configure and return the Gemini generative model."""
        genai.configure(api_key=self.api_key)

        generation_config = {
            "temperature": 0.3,
            "top_p": 0.85,
            "top_k": 40,
            "max_output_tokens": 8192,
        }

        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT",        "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH",       "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]

        return genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            generation_config=generation_config,
            safety_settings=safety_settings,
            system_instruction=self.SYSTEM_PROMPT,
        )

    # ------------------------------------------------------------------
    # Chat session management
    # ------------------------------------------------------------------

    def start_chat(self) -> None:
        """Start a new chat session."""
        self.chat_session = self.model.start_chat(history=[])

    def reset_chat(self) -> None:
        """Reset the current chat session."""
        self.chat_session = None
        self.start_chat()

    # ------------------------------------------------------------------
    # Document processing
    # ------------------------------------------------------------------

    def process_pdf(self, pdf_file) -> ProcessedDocument:
        """Extract text content from a PDF file."""
        try:
            reader = PyPDF2.PdfReader(pdf_file)
            text = "".join(page.extract_text() + "\n" for page in reader.pages)
            return ProcessedDocument(
                name=getattr(pdf_file, 'name', "document.pdf"),
                doc_type="PDF",
                content=text,
                processed_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            )
        except Exception as e:
            raise Exception(f"Error extracting PDF: {e}")

    def process_docx(self, docx_file) -> ProcessedDocument:
        """Extract text content from a DOCX file."""
        try:
            doc = docx.Document(docx_file)
            text = "\n".join(p.text for p in doc.paragraphs)
            return ProcessedDocument(
                name=getattr(docx_file, 'name', "document.docx"),
                doc_type="DOCX",
                content=text,
                processed_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            )
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {e}")

    def process_image(self, image_file) -> ProcessedDocument:
        """Load and wrap an image file for multimodal analysis."""
        try:
            image = Image.open(image_file)
            return ProcessedDocument(
                name=getattr(image_file, 'name', "image.jpg"),
                doc_type="IMAGE",
                content=image,
                processed_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            )
        except Exception as e:
            raise Exception(f"Error processing image: {e}")

    # ------------------------------------------------------------------
    # Voice input — Speech-to-Text
    # ------------------------------------------------------------------

    def transcribe_audio(self, audio_bytes: bytes) -> str:
        """Convert WAV audio bytes to text using Google Speech Recognition."""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
                tmp.write(audio_bytes)
                tmp_path = tmp.name

            with sr.AudioFile(tmp_path) as source:
                audio_data = self.recognizer.record(source)
                text = self.recognizer.recognize_google(audio_data)

            os.unlink(tmp_path)
            return text

        except sr.UnknownValueError:
            raise Exception("Could not understand audio")
        except sr.RequestError as e:
            raise Exception(f"Speech recognition service error: {e}")
        except Exception as e:
            raise Exception(f"Error transcribing audio: {e}")

    # ------------------------------------------------------------------
    # Voice output — Text-to-Speech (Coqui TTS)
    # ------------------------------------------------------------------

    def text_to_speech(
        self,
        text: str,
        language: str = 'english',
        speaker_wav: Optional[str] = None,
        speaker: Optional[str] = None,
    ) -> TTSResult:
        """
        Convert text to speech using Coqui TTS.

        Returns TTSResult which includes:
          - audio_bytes : raw WAV bytes (save to file or stream)
          - audio_b64   : base64-encoded WAV string (embed in HTML <audio> tag directly)
        """
        if self.tts_engine is None:
            raise RuntimeError(
                "TTS is disabled. Initialize LegalAgent with enable_tts=True."
            )
        return self.tts_engine.synthesize(
            text=text,
            language=language,
            speaker_wav=speaker_wav,
            speaker=speaker,
        )

    def clone_voice_and_speak(
        self,
        text: str,
        speaker_wav_path: str,
        language: str = 'english',
    ) -> TTSResult:
        """Generate speech in a cloned voice using XTTS v2."""
        if self.tts_engine is None:
            raise RuntimeError(
                "TTS is disabled. Initialize LegalAgent with enable_tts=True."
            )
        return self.tts_engine.clone_voice_and_speak(
            text=text,
            speaker_wav_path=speaker_wav_path,
            language=language,
        )

    def get_tts_speakers(self, language: str = 'english') -> List[str]:
        """Return available named speakers for the active TTS model."""
        if self.tts_engine is None:
            return []
        return self.tts_engine.list_speakers(language)

    # ------------------------------------------------------------------
    # AI response generation
    # ------------------------------------------------------------------

    def generate_response(
        self,
        user_message: str,
        language: str = 'english',
        documents: Optional[List[ProcessedDocument]] = None,
    ) -> str:
        """Generate an AI text response to a user query."""
        if not self.chat_session:
            self.start_chat()

        # Inject language instruction tag
        if language == 'urdu':
            user_message = f"[RESPOND IN URDU] {user_message}"
        elif language == 'roman':
            user_message = f"[RESPOND IN ROMAN URDU AND ENGLISH] {user_message}"

        # Prepend text document context
        if documents:
            context = "\n\n[UPLOADED DOCUMENTS CONTEXT]:\n"
            for doc in documents:
                if doc.doc_type in ('PDF', 'DOCX'):
                    context += f"\n--- {doc.name} ---\n{doc.content[:2000]}\n"
            user_message = context + "\n[USER QUESTION]: " + user_message

        try:
            if documents and any(d.doc_type == 'IMAGE' for d in documents):
                parts = [user_message]
                for doc in documents:
                    if doc.doc_type == 'IMAGE':
                        parts.append(doc.content)
                response = self.chat_session.send_message(parts)
            else:
                response = self.chat_session.send_message(user_message)

            return response.text

        except Exception as e:
            raise Exception(f"Error generating response: {e}")

    def generate_response_with_audio(
        self,
        user_message: str,
        language: str = 'english',
        documents: Optional[List[ProcessedDocument]] = None,
        speaker_wav: Optional[str] = None,
        speaker: Optional[str] = None,
    ) -> dict:
        """
        Generate a text response AND synthesize it to audio in one call.

        Returns:
            {
                "text":     str        – AI text response,
                "tts":      TTSResult  – synthesized audio with audio_b64 field,
                                         (None if TTS disabled or synthesis failed)
                "language": str
            }
        """
        text_response = self.generate_response(
            user_message=user_message,
            language=language,
            documents=documents,
        )

        tts_result = None
        if self.tts_engine is not None:
            try:
                tts_result = self.text_to_speech(
                    text=text_response,
                    language=language,
                    speaker_wav=speaker_wav,
                    speaker=speaker,
                )
            except Exception as tts_err:
                # TTS failure must never block the text response
                print(f"[TTS WARNING] Audio synthesis failed: {tts_err}")

        return {"text": text_response, "tts": tts_result, "language": language}

    # ------------------------------------------------------------------
    # UI helpers
    # ------------------------------------------------------------------

    def get_welcome_message(self, language: str = 'english') -> str:
        """Return a welcome message in the specified language."""
        if language == 'urdu':
            return """السلام علیکم! PakJurist میں خوش آمدید۔

میں آپ کی مدد کر سکتا ہوں:
- پاکستان کا آئین
- فوجداری قانون
- دیوانی قانون
- خاندانی قانون
- کاروباری قانون

آج میں آپ کی کیسے مدد کر سکتا ہوں؟"""

        return """السلام علیکم! Welcome to PakJurist — Pakistan Legal Awareness System.

I can help you with information about:
- Constitution of Pakistan (Articles, Amendments)
- Criminal Law (PPC, CrPC, FIR procedures)
- Civil Law (Contracts, Property, Torts)
- Family Law (Marriage, Divorce, Inheritance)
- Business & Commercial Law
- Administrative & Regulatory Law

FEATURES:
- Upload PDFs, Documents & Images
- Voice Input  (Speech-to-Text)
- Voice Output (Text-to-Speech via Coqui TTS — multilingual, voice cloning)

How may I assist you today?"""

    def get_quick_questions(self) -> List[str]:
        """Return suggested quick-start questions."""
        return [
            "What are fundamental rights in Pakistan Constitution?",
            "Explain Article 25 of Constitution",
            "What is the FIR process in Pakistan?",
            "Tell me about tenant rights",
            "What is the legal age of marriage?",
        ]

    def get_legal_topics(self) -> List[str]:
        """Return the list of supported legal topic categories."""
        return [
            "Constitution of Pakistan",
            "Criminal Law",
            "Civil Law",
            "Family Law",
            "Business & Commercial Law",
            "Property Law",
            "Labor Law",
            "Tax Law",
        ]
    